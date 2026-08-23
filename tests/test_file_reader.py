from docx import Document
from reportlab.pdfgen import canvas
import pytest

from services.file_reader import read_file


def test_read_txt(tmp_path):
    file_path = tmp_path / "test.txt"

    file_path.write_text(
        "Hello world!",
        encoding="utf-8",
    )

    result = read_file(str(file_path))

    assert result == "Hello world!"


def test_read_txt_persian(tmp_path):
    file_path = tmp_path / "persian.txt"

    file_path.write_text(
        "سلام دنیا!",
        encoding="utf-8",
    )

    result = read_file(str(file_path))

    assert result == "سلام دنیا!"


def test_read_empty_txt(tmp_path):
    file_path = tmp_path / "empty.txt"

    file_path.write_text(
        "",
        encoding="utf-8",
    )

    result = read_file(str(file_path))

    assert result == ""


def test_read_txt_uppercase_extension(tmp_path):
    file_path = tmp_path / "test.TXT"

    file_path.write_text(
        "Python is great.",
        encoding="utf-8",
    )

    result = read_file(str(file_path))

    assert result == "Python is great."


def test_read_pdf(tmp_path):
    file_path = tmp_path / "test.pdf"

    pdf = canvas.Canvas(str(file_path))
    pdf.drawString(100, 750, "Hello PDF")
    pdf.save()

    result = read_file(str(file_path))

    assert "Hello PDF" in result


def test_read_empty_pdf(tmp_path):
    file_path = tmp_path / "empty.pdf"

    pdf = canvas.Canvas(str(file_path))
    pdf.save()

    result = read_file(str(file_path))

    assert result == ""


def test_read_docx(tmp_path):
    file_path = tmp_path / "test.docx"

    document = Document()
    document.add_paragraph("Hello DOCX")
    document.save(str(file_path))

    result = read_file(str(file_path))

    assert result == "Hello DOCX"


def test_read_docx_multiple_paragraphs(tmp_path):
    file_path = tmp_path / "test.docx"

    document = Document()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Second paragraph.")
    document.save(str(file_path))

    result = read_file(str(file_path))

    assert result == (
        "First paragraph.\n"
        "Second paragraph."
    )


def test_read_empty_docx(tmp_path):
    file_path = tmp_path / "empty.docx"

    document = Document()
    document.save(str(file_path))

    result = read_file(str(file_path))

    assert result == ""


def test_read_nonexistent_file(tmp_path):
    file_path = tmp_path / "missing.txt"

    with pytest.raises(FileNotFoundError):
        read_file(str(file_path))


def test_read_unsupported_extension(tmp_path):
    file_path = tmp_path / "test.csv"

    file_path.write_text(
        "hello,world",
        encoding="utf-8",
    )

    with pytest.raises(ValueError):
        read_file(str(file_path))


def test_read_file_with_multiple_lines(tmp_path):
    file_path = tmp_path / "test.txt"

    file_path.write_text(
        "Line one.\n"
        "Line two.\n"
        "Line three.",
        encoding="utf-8",
    )

    result = read_file(str(file_path))

    assert result == (
        "Line one.\n"
        "Line two.\n"
        "Line three."
    )
def test_read_corrupted_pdf(tmp_path):
    file_path = tmp_path / "broken.pdf"

    file_path.write_text(
        "This is not a real PDF file.",
        encoding="utf-8",
    )

    with pytest.raises(Exception):
        read_file(str(file_path))


def test_read_corrupted_docx(tmp_path):
    file_path = tmp_path / "broken.docx"

    file_path.write_text(
        "This is not a real DOCX file.",
        encoding="utf-8",
    )

    with pytest.raises(Exception):
        read_file(str(file_path))


def test_read_directory_instead_of_file(tmp_path):
    directory = tmp_path / "folder"
    directory.mkdir()

    with pytest.raises((IsADirectoryError, PermissionError, OSError)):
        read_file(str(directory))
from docx import Document
from reportlab.pdfgen import canvas

from services.document_builder import build_document


def test_build_document_from_pdf(tmp_path):
    file_path = tmp_path / "test.pdf"

    pdf = canvas.Canvas(str(file_path))
    pdf.drawString(
        100,
        750,
        "Python is powerful.",
    )
    pdf.drawString(
        100,
        730,
        "Python is easy.",
    )
    pdf.save()

    document = build_document(str(file_path))

    assert document.metadata["filename"] == "test.pdf"
    assert document.metadata["extension"] == ".pdf"

    assert document.language == "english"

    assert len(document.sentences) == 2

    assert "python" in document.words
    assert "powerful" in document.words
    assert "easy" in document.words

    assert document.statistics["word_count"] == 6


def test_build_document_from_docx(tmp_path):
    file_path = tmp_path / "test.docx"

    doc = Document()

    doc.add_paragraph(
        "Python is powerful."
    )

    doc.add_paragraph(
        "Python is easy."
    )

    doc.save(str(file_path))

    document = build_document(str(file_path))

    assert document.metadata["filename"] == "test.docx"
    assert document.metadata["extension"] == ".docx"

    assert document.language == "english"

    assert len(document.sentences) == 2

    assert "python" in document.words
    assert "powerful" in document.words
    assert "easy" in document.words

    assert document.statistics["word_count"] == 6


def test_build_document_persian_pdf(tmp_path):
    file_path = tmp_path / "persian.pdf"

    pdf = canvas.Canvas(str(file_path))

    pdf.drawString(
        100,
        750,
        "Salam Donya.",
    )

    pdf.save()

    document = build_document(str(file_path))

    assert document.metadata["extension"] == ".pdf"
    assert document.language == "english"


def test_build_document_empty_docx(tmp_path):
    file_path = tmp_path / "empty.docx"

    doc = Document()
    doc.save(str(file_path))

    document = build_document(str(file_path))

    assert document.normalized_text == ""
    assert document.words == []
    assert document.sentences == []
    assert document.top_words == []

    assert document.statistics["word_count"] == 0